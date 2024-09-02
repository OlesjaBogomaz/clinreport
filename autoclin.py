#! /usr/bin/env python3

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date
from math import log, floor
import argparse
import os
import sqlite3


def main():
    argparser = argparse.ArgumentParser(
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
        description='Clinical report generator',
    )
    argparser.add_argument('-i', '--input-sqlite', required=True, help='path to SQLite')
    argparser.add_argument('-t', '--target-sample', default=False, help='target (proband) sample id')
    argparser.add_argument('-o', '--output-dir', default='.', help='output directory')
    args = argparser.parse_args()

    with sqlite3.connect(args.input_sqlite) as con:
        cur = con.cursor()
        all_samples = [row[0] for row in cur.execute('select distinct base__sample_id from sample;').fetchall()]
        variant_cols = cur.execute('pragma table_info(variant);').fetchall()
        variant_cols = [col[1] for col in variant_cols]
        if 'vep_csq__symbol' not in variant_cols:
            # legacy SQLite
            variant_rows = cur.execute('select * from variant where base__note in (1,2,3);').fetchall()
            variants_data = [dict(zip(variant_cols, row)) for row in variant_rows]
            for varaint_data in variants_data:
                varaint_data.update(annotate_legacy(varaint_data))
        else:
            # new SQLite
            variant_rows = cur.execute('select * from variant where base__note is not null;').fetchall()
            variants_data = [dict(zip(variant_cols, row)) for row in variant_rows]

    for sample in all_samples:
        create_doc(variants_data, sample, all_samples, sample==args.target_sample).save(os.path.join(args.output_dir, f'Заключение ({sample}).docx'))


def create_doc(variants_data: list, sample: str, all_samples: list, target_sample: bool=False, dzm: bool=True) -> Document:
    case_table_data = [(sample, '_', '_', '_')]
    tech_table_data = [(
        'полногеномное секвенирование (Whole Genome Sequencing)',
        '_x',
        'не менее 90 млрд',
        'парно-концевое',
        '150',
        '1.    число прочтений с качеством Q20: не менее 90% от числа прочтений, полученных в результате секвенирования',
        '2.    число прочтений с качеством Q30: не менее 80% от числа прочтений, полученных в результате секвенирования'
    )]
    SNV_P_table_data = form_snv_table_data(filter_variants(variants_data, by_note='1', by_sample=sample))
    SNV_LP_table_data = form_snv_table_data(filter_variants(variants_data, by_note='2', by_sample=sample))
    SNV_VUS_table_data = form_snv_table_data(filter_variants(variants_data, by_note='3', by_sample=sample))
    CNV_table_data = []
    MT_table_data = []
    STR_table_data = []
    SF_table_data = []
    C_table_data = form_snv_table_data(filter_variants(variants_data, by_note='8', by_sample=sample), pathogenicity_col=True)
    variants_data_for_interpretation = sum([filter_variants(variants_data, note, by_sample=sample) for note in ['1', '2', '3']], [])

    doc = Document()
    doc.add_heading('ОТЧЕТ\n', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('по результатам анализа\nданных секвенирования ДНК\n\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(doc, case_table_data, case_table_header, transpose=True)

    doc.add_heading('РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ\n', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Патогенные варианты нуклеотидной последовательности, являющиеся вероятной причиной заболевания', style='List Bullet')
    add_table(doc, SNV_P_table_data, SNV_table_header, italic=True)

    doc.add_paragraph('Вероятно патогенные варианты нуклеотидной последовательности, являющиеся возможной причиной заболевания', style='List Bullet')
    add_table(doc, SNV_LP_table_data, SNV_table_header, italic=True)

    doc.add_paragraph('Варианты нуклеотидной последовательности с неопределенной клинической значимостью', style='List Bullet')
    add_table(doc, SNV_VUS_table_data, SNV_table_header, italic=True)

    doc.add_paragraph('Структурные генетические варианты', style='List Bullet')
    add_table(doc, CNV_table_data, CNV_table_header)

    doc.add_paragraph('Варианты в митохондриальной ДНК', style='List Bullet')
    add_table(doc, MT_table_data, MT_table_header, italic=True)

    doc.add_paragraph('Исследование числа клинически значимых коротких тандемных повторов', style='List Bullet')
    add_table(doc, STR_table_data, STR_table_header, italic=True)

    doc.add_paragraph('Клинически значимые варианты, не связанные с основным диагнозом', style='List Bullet')
    add_table(doc, SF_table_data, SNV_table_header, italic=True)

    if dzm and not target_sample:
        doc.add_paragraph('Носительство вероятно патогенных вариантов, не связанных с основным диагнозом', style='List Bullet')
        add_table(doc, C_table_data, C_table_header)

    doc.add_paragraph('* Частоты аллелей отражают максимальную частоту в популяции и приведены по базе gnomAD v4.1.0 (выборка до 807,162 человек).\n')

    doc.add_heading('ИНТЕРПРЕТАЦИЯ\n', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Был проведен поиск вариантов, ассоциированных с направительным диагнозом у пробанда и прочими наследственными заболеваниями со сходными фенотипическими проявлениями.')

    if not variants_data_for_interpretation:
        doc.add_paragraph('\nЗначимых изменений, соответствующих критериям поиска, не обнаружено.')
    else:
        for variant in variants_data_for_interpretation:
            symbol = variant["vep_csq__symbol"]
            transcript = variant['vep_csq__transcript']
            refseq = variant['vep_csq__refseq']
            hgvsg = variant["vep_csq__hgvsg"]
            hgvsc = variant["vep_csq__hgvsc"]
            hgvsc_msg = f'{refseq}:{hgvsc}' if refseq else f'{transcript}:{hgvsc}'
            hgvsp = variant["vep_csq__hgvsp"]
            hgvsp_msg = f'p.({hgvsp[2:]})' if hgvsp else ''
            rsid = variant['dbsnp__rsid']
            variation_msg = ', '.join([msg for msg in [hgvsg, hgvsc_msg, rsid] if msg])
            consequence = variant["vep_csq__consequence"]
            exon, intron = variant["vep_csq__exon"], variant["vep_csq__intron"]
            if exon:
                gene_part_msg = f"в {exon.split('/')[0]} экзоне из {exon.split('/')[1]} экзонов"
            elif intron:
                gene_part_msg = f"в {intron.split('/')[0]} интроне из {intron.split('/')[1]} интронов"
            if 'missense' in consequence:
                leading_to_msg = f'который приводит к аминокислотной замене {hgvsp_msg}'
            elif 'synon' in consequence:
                leading_to_msg = f'который приводит / может приводить к абберантному сплайсингу {hgvsp_msg}'
            elif 'intron' in consequence:
                leading_to_msg = f'который приводит / может приводить к абберантному сплайсингу'
            elif 'shift' in consequence:
                leading_to_msg = f'который приводит к сдвигу рамки считывания и образованию преждевременного стоп-кодона {hgvsp_msg}'
            elif 'stop' in consequence:
                leading_to_msg = f'который приводит к образованию преждевременного стоп-кодона {hgvsp_msg}'
            elif 'splice' in consequence:
                leading_to_msg = f'который приводит к разрушению канонического сайта сплайсинга'
            omim_pheno, omim_id = variant["vep_omim_pheno__pheno"], variant["vep_omim_pheno__id"]
            gnomad_af = variant["gnomad4genomes__AF"]
            gnomad_af_msg = float2percent(gnomad_af) if gnomad_af else '-'
            zygosity = variant["tagsampler_new__zygosity"]
            zygosity_msg = zygosity2msg[zygosity][2] if zygosity else ''
            ad = variant["tagsampler_new__ad"] or '_'
            ad_msg = f"с глубиной прочтения {ad}x"
            gerp_rs_score = variant["gerp__gerp_rs"]
            insilico_prediction = predict_insilico(variant["dbscsnv__ada_score"], variant["metarnn__score"], variant["revel__score"], variant['alphamissense__score'], variant["phylop100__score"])
            clinvar_id = variant["clinvar_new__id"]
            clinvar_sig = variant["clinvar_new__sig"]
            clinvar_sig_subs = variant["clinvar_new__sig_subs"]
            clinvar_sig_subs_msgs = clinvar_sig_subs2msgs(clinvar_sig_subs) or [f'как {clinsig2msg.get(clinvar_sig, clinvar_sig)}']
            clinvar_equivalents = variant["clinvar_new__equivalents"]
            clinvar_alternatives = variant["clinvar_new__alternatives"]
            clinvar_equivalents = eval(clinvar_equivalents) if clinvar_equivalents else []
            clinvar_alternatives = eval(clinvar_alternatives) if clinvar_alternatives else []
            samples = variant["tagsampler_new__samples"]
            samples = samples.split(';') if samples else []
            clinsig = note2clinsig[variant["base__note"]]

            intro_paragraph = doc.add_paragraph('\n')
            intro_paragraph.add_run(f'Обнаружен ранее _ описанный в литературе вариант ({variation_msg}) {zygosity_msg} {gene_part_msg} гена ')
            intro_paragraph.add_run(f'{symbol}').italic = True
            intro_paragraph.add_run(f', {leading_to_msg}, {ad_msg}.')

            if omim_pheno:
                omim_paragraph = doc.add_paragraph()
                omim_paragraph.add_run('Патогенные варианты в гене ')
                omim_paragraph.add_run(f'{symbol}').italic = True
                omim_paragraph.add_run(f' приводят к {omim_pheno} ({omim_id}).')

            if gnomad_af:
                doc.add_paragraph(f'Вариант встречается с частотой {gnomad_af_msg} в базах данных популяционных частот gnomAD')
            else:
                doc.add_paragraph('Вариант не встречается в базах данных популяционных частот gnomAD.')

            comp_paragraph = doc.add_paragraph()
            if 'missense' in consequence:
                if gerp_rs_score:
                    if gerp_rs_score >= 2:
                        comp_paragraph.add_run('Вариант расположен в высококонсервативной позиции. ')
                    elif gerp_rs_score >= 0:
                        comp_paragraph.add_run('Вариант расположен в консервативной позиции. ')
                    else:
                        comp_paragraph.add_run('Вариант расположен в неконсервативной позиции. ')
                if insilico_prediction:
                    comp_paragraph.add_run('Компьютерные алгоритмы предсказывают патогенный эффект варианта на белок.')
                else:
                    comp_paragraph.add_run('Компьютерные алгоритмы предсказывают нейтральный эффект варианта на белок.')
            elif 'shift' in consequence or 'stop' in consequence:
                comp_paragraph.add_run('Вариант с большой долей вероятности приводит к потере функции соответствующей копии гена.')
            elif 'splice' in consequence:
                if insilico_prediction:
                    comp_paragraph.add_run('Вариант предсказан приводить к аберрантному сплайсингу компьютерными алгоритмами. ')
                    comp_paragraph.add_run('Вариант с большой долей вероятности приводит к потере функции соответствующей копии гена.')
                else:
                    comp_paragraph.add_run('Вариант не предсказан приводить к аберрантному сплайсингу компьютерными алгоритмами. ')
            elif 'synon' in consequence or 'intron' in consequence:
                if insilico_prediction:
                    comp_paragraph.add_run('Вариант предсказан приводить к аберрантному сплайсингу компьютерными алгоритмами. ')
                else:
                    comp_paragraph.add_run('Вариант не предсказан приводить к аберрантному сплайсингу компьютерными алгоритмами. ')
                comp_paragraph.add_run('Требуется проведение функционального анализа.')

            if clinvar_sig_subs_msgs:
                sources.append(f'https://www.ncbi.nlm.nih.gov/clinvar/variation/{clinvar_id}')
                source_idx = len(sources)
                doc.add_paragraph(f'Вариант аннотирован {", ".join(clinvar_sig_subs_msgs)} в базе данных ClinVar [{source_idx}].')

            for clinvar_equivalent in clinvar_equivalents:
                sources.append(f'https://www.ncbi.nlm.nih.gov/clinvar/variation/{clinvar_equivalent[0]}')
                source_idx = len(sources)
                clinvar_equivalent_sig_subs_msgs = clinvar_sig_subs2msgs(clinvar_equivalent[4]) or [clinvar_equivalent[2]]
                doc.add_paragraph(f'Вариант с такой же аминокислотной заменой {clinvar_equivalent[1]} в той же позиции аннотирован {", ".join(clinvar_equivalent_sig_subs_msgs)} [{source_idx}].')

            for clinvar_alternative in clinvar_alternatives:
                sources.append(f'https://www.ncbi.nlm.nih.gov/clinvar/variation/{clinvar_alternative[0]}')
                source_idx = len(sources)
                clinvar_alternative_sig_subs_msgs = clinvar_sig_subs2msgs(clinvar_alternative[4]) or [clinvar_alternative[2]]
                doc.add_paragraph(f'Вариант с другой аминокислотной заменой {clinvar_alternative[1]} в той же позиции аннотирован {", ".join(clinvar_alternative_sig_subs_msgs)} [{source_idx}].')

            if target_sample:
                nontarget_samples = samples.copy()
                nontarget_samples.remove(sample)
                if nontarget_samples:
                    doc.add_paragraph(f'Вариант обнаружен у {", ".join(nontarget_samples)}')
                # else:
                #     all_nontarget_samples = all_samples.copy()
                #     all_nontarget_samples.remove(sample)
                #     doc.add_paragraph(f'Вариант не обнаружен у {", ".join(all_nontarget_samples)}, таким образом является de novo.')

            doc.add_paragraph(f'По совокупности сведений вариант расценивается как {clinsig}.')
            doc.add_paragraph('Рекомендуется сопоставление фенотипа пациента с фенотипом заболеваний, ассоциированных с геном.')
            doc.add_paragraph('Вариант требует обязательного подтверждения генотипа референсным методом (секвенирование по методу Сэнгера).')

        doc.add_paragraph('\nДругих значимых изменений, соответствующих критериям поиска, не обнаружено.')

    doc.add_paragraph('Оценка клинической значимости (патогенности) выявленных вариантов проводилась на основании российских рекомендаций для интерпретации данных, полученных методами массового параллельного секвенирования (MPS).')
    doc.add_paragraph().add_run('Результаты данного исследования могут быть правильно интерпретированы только врачом-генетиком.').bold = True

    doc.add_heading('ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ\n', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, tech_table_data, tech_table_header, transpose=True)

    doc.add_heading('СПИСОК ЛИТЕРАТУРЫ И БАЗ ДАННЫХ\n', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for source in sources:
        doc.add_paragraph(source, style='List Number')
    doc.add_paragraph('\n')

    doc.add_paragraph(f'Дата выдачи отчета: {date.today()}')
    doc.add_paragraph('Клинический биоинформатик: ')

    return doc


def filter_variants(variants_data: list, by_note: str, by_sample=None) -> list:
    variants_data = [variant for variant in variants_data if variant['base__note'] == by_note]
    if by_sample:
        variants_data_filtered = []
        for variant in variants_data:
            variant_copy = variant.copy()
            if by_sample in variant_copy['tagsampler_new__samples'].split(';'):
                sample_idx = variant_copy['tagsampler_new__samples'].split(';').index(by_sample)
                if variant_copy['tagsampler_new__filter'].split(';')[sample_idx] != 'PASS':
                    continue
                variant_copy['tagsampler_new__zygosity'] = variant_copy['tagsampler_new__zygosity'].split(';')[sample_idx]
                variant_copy['tagsampler_new__ad'] = variant_copy['tagsampler_new__ad'].split(';')[sample_idx].split(',')[-1]
                variant_copy['tagsampler_new__dp'] = variant_copy['tagsampler_new__dp'].split(';')[sample_idx]
                variants_data_filtered.append(variant_copy)
    else:
        variants_data_filtered = variants_data.copy()
    return variants_data_filtered


def form_snv_table_data(variants_data: list, pathogenicity_col=False) -> list:
    snv_table_data = []
    for variant in variants_data:
        symbol = variant['vep_csq__symbol']
        omim_pheno = variant['vep_omim_pheno__pheno']
        chrom = variant['base__chrom']
        pos = variant['extra_vcf_info__pos']
        ref = variant['extra_vcf_info__ref']
        alt = variant['extra_vcf_info__alt']
        spdi = f'{chrom}-{pos}-{ref}-{alt}'
        rsid = variant['dbsnp__rsid'] or ''
        hgvsc = variant['vep_csq__hgvsc']
        hgvsp = variant['vep_csq__hgvsp']
        hgvsp_msg = f' p.({hgvsp[2:]})' if hgvsp else ''
        transcript = variant['vep_csq__transcript']
        refseq = variant['vep_csq__refseq']
        hgvsc_msg = f'{refseq}:{hgvsc}' if refseq else f'{transcript}:{hgvsc}'
        variation = '\n'.join([msg for msg in [spdi, hgvsc_msg, hgvsp_msg, rsid] if msg])
        zygosity = variant['tagsampler_new__zygosity']
        zygosity_msg = zygosity2msg[zygosity][1] if zygosity else '-'
        inher = variant['vep_omim_pheno__inher']
        inher_msg = ', '.join(inher2msg[inh] for inh in inher.split(',')) if inher else '-'
        zyg_inher_msg = f'{zygosity_msg}\n({inher_msg})'
        af = variant['gnomad4genomes__AF']
        af_msg = float2percent(af) if af else 'н/д'
        ad = variant['tagsampler_new__ad'] or '_'
        dp = variant['tagsampler_new__dp'] or '_'
        cover_msg = f'{ad}x/{dp}x'
        if pathogenicity_col:
            clinvar_sig = variant["clinvar_new__sig"]
            clinsig_msg = clinsig2msg.get(clinvar_sig, '-')
            snv_table_data.append((symbol, omim_pheno, variation, zyg_inher_msg, clinsig_msg, af_msg, cover_msg))
        else:
            snv_table_data.append((symbol, omim_pheno, variation, zyg_inher_msg, af_msg, cover_msg))
    return snv_table_data


def add_table(document: Document, table_data: list, table_header: tuple, italic: bool=False, transpose: bool=False) -> None:
    table_data.insert(0, table_header)
    if transpose:
        table_data = list(zip(*table_data))
    table = document.add_table(rows=len(table_data), cols=len(table_data[0]))
    for i in range(len(table_data)):
        for j in range(len(table_data[0])):
            cell_paragraph_run = table.rows[i].cells[j].add_paragraph().add_run(str(table_data[i][j]))
            if italic and i > 0 and j == 0:
                cell_paragraph_run.italic = True
    if len(table_data) == 1:
        table.add_row()
        table.cell(i+1, 0).merge(table.cell(i+1, j))
        table.rows[i+1].cells[0].text = 'Не обнаружено'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'


def float2percent(f: float) -> str:
    ndigits = max(1, -floor(log(100*f, 10)))
    if round(10**ndigits * 100*f) == 1:
        ndigits += 1
    return f'{round(100*f, ndigits)}%'


def predict_insilico(scsnvada, metarnn, revel, alphamissense, phylop):
    thresholds = {
        'scsnvada': (0.957813, 0.999322, 0.999925),
        'metarnn': (0.748, 0.841, 0.939),
        'revel': (0.644, 0.773, 0.932),
        'alphamissense': (0.787, 0.956, 0.994),
        'phylop': (7.52, 9.88, float('inf'))
    }
    if scsnvada:
        return True if scsnvada >= thresholds['scsnvada'][0] else False
    if metarnn:
        return True if metarnn >= thresholds['metarnn'][0] else False
    if revel:
        return True if revel >= thresholds['revel'][0] else False
    if alphamissense:
        return True if alphamissense >= thresholds['alphamissense'][0] else False
    if phylop:
        return True if phylop >= thresholds['phylop'][0] else False
    return False


def clinvar_sig_subs2msgs(clinvar_sig_subs) -> list:
    sig_subs_msgs = []
    if not clinvar_sig_subs:
        return sig_subs_msgs
    for sig_count in clinvar_sig_subs.split('; '):
        sig, count = sig_count[:-1].split(' (')
        sig_subs_msgs.append(f'как {clinsig2msg.get(sig, sig)} {count} лабораторией(ями)')
    return sig_subs_msgs


def annotate_legacy(variant_data: dict) -> dict:
    extra_vcf_info = get_extra_vcf_info(variant_data)
    for i in range(extra_vcf_info['nblocks']):
        if not extra_vcf_info['CSQ_PICK'][i] == '1':
            continue
        annotation = {
            'vep_csq__symbol': extra_vcf_info['CSQ_SYMBOL'][i],
            'vep_csq__transcript': extra_vcf_info['CSQ_Feature'][i],
            'vep_csq__hgvsc': extra_vcf_info['CSQ_HGVSc'][i].split(':')[-1],
            'vep_csq__hgvsp': extra_vcf_info['CSQ_HGVSp'][i].split(':')[-1],
            'vep_csq__hgvsg': extra_vcf_info['CSQ_HGVSg'][i],
            'vep_csq__consequence': extra_vcf_info['CSQ_Consequence'][i],
            'vep_csq__biotype': extra_vcf_info['CSQ_BIOTYPE'][i],
            'vep_csq__exon': extra_vcf_info['CSQ_EXON'][i],
            'vep_csq__intron': extra_vcf_info['CSQ_INTRON'][i],
            'vep_csq__strand': extra_vcf_info['CSQ_STRAND'][i],
            'vep_csq__codons': extra_vcf_info['CSQ_Codons'][i],
        }
        annotation['vep_csq__refseq'] = extra_vcf_info['CSQ_MANE_SELECT'][i] if extra_vcf_info['CSQ_MANE_SELECT'][i] else None
    annotation['vep_omim_pheno__inher'] = get_inher_from_omim_pheno(variant_data['vep_omim_pheno__pheno'])
    for col in ['filter', 'zygosity', 'ad', 'dp']:
        annotation[f'tagsampler_new__{col}'] = variant_data[f'vevatacmg_postaggregator__{col}']
    annotation['tagsampler_new__samples'] = variant_data[f'vevatacmg_postaggregator__sample']
    for col in ['id', 'sig']:
        annotation[f'clinvar_new__{col}'] = variant_data[f'clinvar__{col}']
    annotation['clinvar_new__sig_subs'] = annotation['clinvar_new__equivalents'] = annotation['clinvar_new__alternatives'] = None
    return annotation


def get_extra_vcf_info(variant_data: dict) -> dict:
    """
    Make each CSQ block iterable
    """
    nblocks = len(variant_data['extra_vcf_info__CSQ_Allele'].split(';'))
    variant_data_transformed = {'nblocks': nblocks}
    for key, value in variant_data.items():
        if key.startswith('extra_vcf_info__CSQ'):
            if value is None:
                value = ['']*nblocks
            else:
                value = value.split(';')
        variant_data_transformed[key.lstrip('extra_vcf_info__')] = value
    return variant_data_transformed


def get_inher_from_omim_pheno(phenotype: str) -> str:
    if not phenotype:
        return None
    inheritance_map = {
        'Autosomal dominant': 'AD',
        'X-linked dominant': 'XD',
        'Autosomal recessive': 'AR',
        'X-linked recessive': 'XR'
    }
    inher = set()
    for name, short in inheritance_map.items():
        if name in phenotype:
            inher.add(short)
    inher = ','.join(sorted(inher))
    return inher


note2clinsig = {
    '1': 'патогенный',
    '2': 'вероятно патогенный',
    '3': 'вариант с неизвестной клинической значимостью'
}
clinsig2msg = {
    'Pathogenic': 'патогенный',
    'Pathogenic/Likely_pathogenic': 'патогенный / вероятно патогенный',
    'Pathogenic/Likely pathogenic': 'патогенный / вероятно патогенный',
    'Likely_pathogenic': 'вероятно патогенный',
    'Likely pathogenic': 'вероятно патогенный',
    'Uncertain_significance': 'вариант с неизвестной клинической значимостью',
    'Uncertain significance': 'вариант с неизвестной клинической значимостью',
}
zygosity2msg = {
    'het': {1: 'Гетерозигота', 2: 'в гетерозиготном состоянии'},
    'hom': {1: 'Гомозигота', 2: 'в гомозиготном состоянии'}
}
inher2msg = {
    'AD': 'Аутосомно-доминантный',
    'XD': 'Х-сцепленный доминантный',
    'AR': 'Аутосомно-рецессивный',
    'XR': 'Х-сцепленный рецессивный'
}
case_table_header = (
    'Номер образца',
    'Пол пациента',
    'Возраст пациента',
    'Предварительный диагноз'
)
tech_table_header = (
    'Метод исследования',
    'Средняя глубина прочтения генома после секвенирования',
    'Количество прочитанных нуклеотидов',
    'Тип прочтения',
    'Длина прочтения',
    'Качество выходных данных секвенирования',
    ''
)
SNV_table_header = (
    'Ген',
    'Ассоциированное заболевание (OMIM)',
    'Изменение ДНК (HG38) (Изменение белка)',
    'Зиготность (Тип наследования)',
    'Частота*',
    'Кол-во прочтений (АЛТ/ОБЩ)'
)
CNV_table_header = (
    'Изменение ДНК (HG38)',
    'Ассоциированное заболевание (OMIM)',
    'Затронутые морбидные гены',
    'Число копий',
    'Классификация'
)
MT_table_header = (
    'Ген',
    'Ассоциированное заболевание (OMIM)',
    'Изменение ДНК',
    'Классификация',
    'Кол-во прочтений (АЛТ/ОБЩ)'
)
STR_table_header = (
    'Ген',
    'Ассоциированное заболевание (OMIM)',
    'Экспансия повтора',
    'Оценочное число повторов',
    'Классификация'
)
C_table_header = (
    'Ген',
    'Ассоциированное заболевание (OMIM)',
    'Изменение ДНК (HG38) (Изменение белка)',
    'Зиготность (Тип наследования)',
    'Патогенность',
    'Частота*',
    'Кол-во прочтений (АЛТ/ОБЩ)'
)
sources = [
    'http://www.omim.org/',
    'http://www.ncbi.nlm.nih.gov/snp/',
    'http://gnomad.broadinstitute.org/',
    'http://www.ncbi.nlm.nih.gov/clinvar/',
    'https://franklin.genoox.com/clinical',
    'http://www.ensembl.org/',
    'https://www.uniprot.org/',
    'https://www.deciphergenomics.org/',
    'https://mitomap.org',
    'https://www.clinicalgenome.org/'
]


if __name__ == '__main__':
    main()